import json
import shutil
import pandas as pd
import os
import sys
import pickle

from lekit.Str.Core                     import UnWrapper, list_byte_to_string
from lekit.Lang                         import BaseClass

from typing                                         import *
from pathlib                                        import Path
from pydub                                          import AudioSegment
from PIL                                            import Image, ImageFile
from docx                                           import Document
from docx.document                                  import Document as DocumentObject

audio_file_type = ["mp3","ogg","wav"]
image_file_type = ['png', 'jpg', 'jpeg', 'bmp', 'svg', 'ico']
temp_tool_file_path_name = "temp.tool_file"

is_binary_file_functional_test_length:int = 1024

def is_binary_file(file_path: str) -> bool:
    try:
        global is_binary_file_functional_test_length
        with open(file_path, 'rb') as f:
            chunk = f.read(is_binary_file_functional_test_length)  # 读取文件的前缀字节
            return b'\x00' in chunk  # 如果包含NUL字符，则认为是二进制文件
    except Exception as e:
        print(f"Error: {e}")
        return False

def get_extension_name(file:str):
        return os.path.splitext(file)[1][1:]

def get_base_filename(file:str):
    return os.path.basename(file)

def is_image_file(file_path:str):
    return get_extension_name(file_path) in image_file_type

dir_name_type = str
file_name_type = str

class tool_file:

    __datas_lit_key:    Literal["model"] = "model"

    def __init__(self, file_path:str, file_mode:str=None, *args, **kwargs):
        self.__file_path:   str             = file_path
        self.datas:         Dict[str, Any]  = {}
        self.__file:        IO[Any]         = None
        if file_mode is not None:
            self.open(file_mode, *args, **kwargs)
    def __del__(self):
        self.close()
    def __str__(self):
        return self.get_path()
    def __setitem__(self, key:str, value):
        self.datas[key] = value
    def __getitem__(self, key:str):
        return self.datas[key]
    def __contains__(self, key:str):
        return key in self.datas
    def __delitem__(self, key:str):
        del self.datas[key]
    def __iter__(self):
        return iter(self.datas)
    def __len__(self):
        return len(self.datas)

    def __or__(self, other):
        if other is None:
            return tool_file(self.get_path() if self.is_dir() else self.get_path()+"\\")
        else:
            return tool_file(os.path.join(self.get_path(), UnWrapper(other)))
    def __idiv__(self, other):
        self.close()
        temp = self.__or__(other)
        self.__file_path = temp.get_path()

    def to_path(self):
        return Path(self.__file_path)
    def __Path__(self):
        return Path(self.__file_path)

    def create(self):
        if self.exists() == False:
            if self.is_dir():
                if os.path.exists(self.get_dir()):
                    os.makedirs(self.__file_path)
                else:
                    raise FileNotFoundError(f"{self.__file_path} cannt create, because its parent path is not exist")
            else:
                self.open('w')
                self.close()
        return self
    def exists(self):
        return os.path.exists(self.__file_path)
    def remove(self):
        self.close()
        if self.exists():
            if self.is_dir():
                shutil.rmtree(self.__file_path)
            else:
                os.remove(self.__file_path)
        return self
    def copy(self, to_path:Union[Self, str]):
        if self.exists() is False:
            raise FileNotFoundError("file not found")
        self.close()
        target_file = tool_file(UnWrapper(to_path))
        if target_file.is_dir():
            target_file = target_file|self.get_filename()
        shutil.copy(self.__file_path, UnWrapper(target_file))
        return target_file
    def move(self, to_path:Union[Self, str]):
        if self.exists() is False:
            raise FileNotFoundError("file not found")
        self.close()
        target_file = tool_file(UnWrapper(to_path))
        if target_file.is_dir():
            target_file = target_file|self.get_filename()
        shutil.move(self.__file_path, UnWrapper(target_file))
        self.__file_path = target_file.get_path()
        return self
    def rename(self, newpath:Union[Self, str]):
        if self.exists() is False:
            raise FileNotFoundError("file not found")
        self.close()
        newpath:str = UnWrapper(newpath)
        if '\\' in newpath or '/' in newpath:
            newpath = get_base_filename(newpath)
        new_current_path = os.path.join(self.get_dir(), newpath)
        os.rename(self.__file_path, new_current_path)
        self.__file_path = new_current_path
        return self

    def refresh(self):
        self.load()
        return self
    def open(self, mode='r', is_refresh=False, encoding:str='utf-8', *args, **kwargs):
        self.close()
        self.__file = open(self.__file_path, mode, encoding=encoding, *args, **kwargs)
        if is_refresh:
            self.refresh()
        return self.__file
    def close(self):
        if self.__file:
            self.__file.close()
        if self.__datas_lit_key in self.datas:
            self.datas[self.__datas_lit_key] = None
        return self.__file
    def is_open(self)->bool:
        return self.__file

    def load(self):
        if self.__file_path is None:
            raise FileNotFoundError("file path is none")
        elif self.is_dir():
            self.__file = open(os.path.join(self.get_path(), temp_tool_file_path_name), 'rb')
            self.data = pickle.load(self.__file)
            return self.data
        suffix = self.get_extension()
        if suffix == 'json':
            self.load_as_json()
        elif suffix == 'csv':
            self.load_as_csv()
        elif suffix == 'xml':
            self.load_as_xml()
        elif suffix == 'xlsx' or suffix == 'xls':
            self.load_as_excel()
        elif suffix == 'txt':
            self.load_as_text()
        elif suffix == 'docx':
            self.load_as_docx()
        elif suffix in audio_file_type:
            self.load_as_audio()
        elif is_binary_file(self.__file_path):
            self.load_as_binary()
        elif is_image_file(self.__file_path):
            self.load_as_image()
        else:
            self.load_as_text()
        return self.data
    def load_as_json(self) -> pd.DataFrame:
        self.open('r')
        self.data = json.load(self.__file)
        return self.data
    def load_as_csv(self) -> pd.DataFrame:
        self.open('r')
        self.data = pd.read_csv(self.__file)
        return self.data
    def load_as_xml(self) -> pd.DataFrame:
        self.open('r')
        self.data = pd.read_xml(self.__file)
        return self.data
    def load_as_dataframe(self) -> pd.DataFrame:
        self.open('r')
        self.data = pd.read_csv(self.__file)
        return self.data
    def load_as_excel(self) -> pd.DataFrame:
        self.open('r')
        self.data = pd.read_excel(self.__file)
        return self.data
    def load_as_binary(self) -> bytes:
        self.open('rb')
        self.data = self.__file.read()
        return self.data
    def load_as_text(self) -> str:
        self.open('r')
        self.data = list_byte_to_string(self.__file.readlines())
        return self.data
    def load_as_wav(self):
        self.data = AudioSegment.from_wav(self.__file_path)
        return self.data
    def load_as_audio(self):
        self.data = AudioSegment.from_file(self.__file_path)
        return self.data
    def load_as_image(self) -> ImageFile.ImageFile:
        self.data = Image.open(self.__file_path)
        return self.data
    def load_as_docx(self) -> DocumentObject:
        self.data = Document(self.__file_path)
        return self.data

    def save(self, path:str=None):
        if path is None and self.__file_path is None:
            raise Exception('No file path specified')
        elif path is None and self.is_dir():
            with open(os.path.join(self.__file_path, temp_tool_file_path_name),'wb') as temp_file:
                pickle.dump(self.data, temp_file)
            return self
        suffix = self.get_extension(path)
        if suffix == 'json':
            self.save_as_json(path)
        elif suffix == 'csv':
            self.save_as_csv(path)
        elif suffix == 'xml':
            self.save_as_xml(path)
        elif suffix == 'xlsx' or suffix == 'xls':
            self.save_as_excel(path)
        elif suffix == 'txt':
            self.save_as_text(path)
        elif suffix == 'docx':
            self.save_as_docx(path)
        elif suffix in audio_file_type:
            self.save_as_audio(path, suffix)
        elif is_binary_file(self.__file_path):
            self.save_as_binary(path)
        elif is_image_file(self.__file_path):
            self.save_as_image(path)
        else:
            self.save_as_text(path)
        return self
    def save_as_json(self, path:str):
        path = path if path else self.__file_path
        json.dump(self.data, self.__file)
        return self
    def save_as_csv(self, path:str):
        path = path if path else self.__file_path
        self.data.to_csv(path)
        return self
    def save_as_xml(self, path:str):
        path = path if path else self.__file_path
        self.data.to_xml(path)
        return self
    def save_as_dataframe(self, path:str):
        path = path if path else self.__file_path
        self.data.to_csv(path)
        return self
    def save_as_excel(self, path:str):
        path = path if path else self.__file_path
        self.data.to_excel(path, index=False)
        return self
    def save_as_binary(self, path:str):
        if path is not None:
            with open(path, 'wb') as f:
                f.write(self.data)
        else:
            self.__file.write(self.data)
        return self
    def save_as_text(self, path:str):
        if path is not None:
            with open(path, 'w') as f:
                f.writelines(self.data)
        else:
            self.__file.writelines(self.data)
        return self
    def save_as_audio(self, path:str):
        self.data.export(path if path else self.__file_path, format=self.get_extension(path))
        return self
    def save_as_image(self, path:str):
        self.data.save(path if path else self.__file_path)
        return self
    def save_as_docx(self, path:str):
        if self.data is str:
            self.data = Document()
            table = self.data.add_table(rows=1, cols=1)
            table.cell(0, 0).text = self.data
        self.data.save(path if path else self.__file_path)
        return self

    def get_size(self) -> int:
        if self.is_dir():
            return -1
        return os.path.getsize(self.__file_path)
    def get_data_type(self) -> type:
        return type(self.data)
    def has_data_type_is(self, types:Union[type, Sequence[type]]) -> bool:
        if isinstance(types, Sequence) is False:
            return self.get_data_type() == types
        return self.get_data_type() in types
    def get_extension(self, path:str=None):
        if self.is_dir() and path is None:
            raise Exception("Cannot get extension of a directory")
        path = path if path is not None else self.__file_path
        if path is None:
            raise Exception("Cannot get extension without target path")
        return get_extension_name(path)
    def get_path(self):
        return self.__file_path
    def get_filename(self, is_without_extension = False):
        '''
        if target path is a file, it return filename
        if target path is a directory, it return top directory name
        '''
        if is_without_extension and '.' in self.__file_path:
            return get_base_filename(self.__file_path)[:-(len(self.get_extension())+1)]
        elif self.__file_path[-1] == '/' or self.__file_path[-1] == '\\':
            return get_base_filename(self.__file_path[:-1])
        else:
            return get_base_filename(self.__file_path)
    def get_dir(self):
        return os.path.dirname(self.__file_path)
    def get_current_dir_name(self):
        return os.path.dirname(self.__file_path)

    def is_dir(self):
        if self.__file_path[-1] == '\\' or self.get_path()[-1] == '/':
            return True
        else:
            return os.path.isdir(self.__file_path)
    def is_file(self):
        return os.path.isfile(self.__file_path)
    def is_binary_file(self):
        return is_binary_file(self.__file)
    def is_image(self):
         return is_image_file(self.__file_path)

    def try_create_parent_path(self):
        dir_path = os.path.dirname(self.__file_path)
        if dir_path == '':
            return self
        if not os.path.exists(dir_path):
            os.makedirs(dir_path)
        return self
    def dir_iter(self):
        return os.listdir(self.__file_path)
    def dir_tool_file_iter(self):
        result = [self]
        result.clear()
        for file in os.listdir(self.__file_path):
            result.append(self|file)
        return result
    def back_to_parent_dir(self):
        self.close()
        self.__file_path = self.get_dir()
        return self
    def dir_count(self, ignore_folder:bool = True):
        iter    = self.dir_iter()
        result  = 0
        for content in iter:
            if ignore_folder and os.path.isdir(os.path.join(self.__file_path, content)):
                continue
            result += 1
    def dir_clear(self):
        for file in self.dir_tool_file_iter():
            file.remove()
        return self
    def first_file_with_extension(self, extension:str):
        target_dir = self if self.is_dir() else tool_file(self.get_dir())
        for file in target_dir.dir_tool_file_iter():
            if file.is_dir() is False and file.get_extension() == extension:
                return file
        return None
    def first_file(self, pr:Callable[[str], bool]):
        target_dir = self if self.is_dir() else tool_file(self.get_dir())
        for file in target_dir.dir_tool_file_iter():
            if pr(file.get_filename()):
                return file
        return None
    def find_file_with_extension(self, extension:str):
        target_dir = self if self.is_dir() else tool_file(self.get_dir())
        result:List[tool_file] = []
        for file in target_dir.dir_tool_file_iter():
            if file.is_dir() is False and file.get_extension() == extension:
                result.append(file)
        return result
    def find_file(self, pr:Callable[[str], bool]):
        target_dir = self if self.is_dir() else tool_file(self.get_dir())
        result:List[tool_file] = []
        for file in target_dir.dir_tool_file_iter():
            if pr(file.get_filename()):
                result.append(file)
        return result
    def dir_walk(
        self,
        top,
        topdown:        bool               = True,
        onerror:        Optional[Callable] = None,
        followlinks:    bool               = False
        ) -> Iterator[tuple[dir_name_type, list[dir_name_type], list[file_name_type]]]:
        return os.walk(self.__file_path)

    def append_text(self, line:str):
        if self.has_data_type_is(type(str)):
            self.data += line
        elif self.has_data_type_is(type(DocumentObject)):
            self.data.add_paragraph(line)
        else:
            raise TypeError(f"Unsupported data type for {sys._getframe().f_code.co_name}")
        return self

    def bool(self):
        return self.exists()
    def __bool__(self):
        return self.exists()

    def must_exists_as_new(self):
        self.close()
        self.try_create_parent_path()
        self.create()
        return self
    def must_exists_path(self):
        return self.must_exists_as_new()

    def make_file_inside(self, data:Self, is_delete_source = False):
        if self.is_dir() is False:
            raise Exception("Cannot make file inside a file, because this object target is not a directory")
        result = self|data.get_filename()
        if is_delete_source:
            data.move(result)
        else:
            data.copy(result)
        return self

    @property
    def extension(self):
        if self.is_dir():
            return None
        return self.get_extension()
    @property
    def filename(self):
        if self.is_dir():
            return None
        return self.get_filename(True)
    @property
    def dirname(self):
        if self.is_dir():
            return self.get_current_dir_name()
        return None
    @property
    def dirpath(self):
        if self.is_dir():
            return self.get_current_dir_name()
        return None
    @property
    def shortname(self):
        return self.get_filename(False)
    @property
    def fullpath(self):
        return self.get_path()

    def make_lib_path(self) -> Path:
        return Path(self.__file_path)

    def in_extensions(self, *args:str) -> bool:
        return self.get_extension() in args

def Wrapper(file) -> tool_file:
    if isinstance(file, tool_file):
        return file
    else:
        return tool_file(UnWrapper(file))

def split_elements(
    file:               Union[tool_file, str],
    *,
    ratios:             List[float]                                 = [1,1],
    pr:                 Optional[Callable[[tool_file], bool]]       = None,
    shuffler:           Optional[Callable[[List[tool_file]], None]] = None,
    output_dirs:        Optional[List[tool_file]]                   = None,
    output_must_exist:  bool                                        = True,
    output_callback:    Optional[Callable[[tool_file], None]]       = None
    ) -> List[List[tool_file]]:
    result:                 List[List[tool_file]]   = BaseClass.split_elements(Wrapper(file).dir_tool_file_iter(),
                                      ratios=ratios,
                                      pr=pr,
                                      shuffler=shuffler)
    if output_dirs is None:
        return result
    for i in range(min(len(output_dirs), len(result))):
        output_dir:         tool_file               = output_dirs[i]
        if output_dir.is_dir() is False:
            raise Exception("Outputs must be directory")
        if output_must_exist:
            output_dir.must_exists_as_new()
        for file in result[i]:
            current = output_dirs[i].make_file_inside(file)
            if output_callback:
                output_callback(current)

    return result

tool_file_or_str = Union[tool_file, str]

if __name__ == "__main__":
    a = tool_file("abc/")
    b = tool_file("a.x")
    c = a|b
    print(c.get_path())


